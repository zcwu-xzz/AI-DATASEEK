import importlib.util
from pathlib import Path

import pymupdf as fitz
from docx import Document


ROOT = Path(__file__).resolve().parents[2]
SPEC = importlib.util.spec_from_file_location("document_operations", ROOT / "tools" / "documents" / "operations.py")
OPS = importlib.util.module_from_spec(SPEC)
assert SPEC.loader is not None
SPEC.loader.exec_module(OPS)


def _pdf(path,text="Temperature observations for 2024"):
    doc=fitz.open(); page=doc.new_page(); page.insert_text((72,72),text); doc.save(path); doc.close()


def _docx(path):
    doc=Document(); doc.add_heading("Dataset description",level=1); doc.add_paragraph("Temperature observations for 2024")
    table=doc.add_table(rows=3,cols=2); table.cell(0,0).text="region"; table.cell(0,1).text="value"; table.cell(1,0).text="A"; table.cell(1,1).text="1"; table.cell(2,0).text="B"; table.cell(2,1).text="2"; doc.save(path)


def test_pdf_inspect_text_render_evidence_and_validate(tmp_path,monkeypatch):
    monkeypatch.setenv("AI_DATASEEK_OUTPUT_ROOT",str(tmp_path)); path=tmp_path/"sample.pdf"; _pdf(path)
    assert OPS.inspect_document({"input_path":str(path)})["summary"]["pages"]==1
    assert "Temperature" in OPS.pdf_text({"input_path":str(path),"max_chars":10000})["summary"]["pages"][0]["text"]
    evidence=OPS.pdf_evidence({"input_path":str(path),"queries":["Temperature"],"max_hits":5,"context_chars":20})
    assert evidence["summary"]["hit_count"]==1
    rendered=OPS.pdf_render({"input_path":str(path),"output_dir":str(tmp_path/"pages"),"max_pages":2,"dpi":72})
    assert rendered["artifacts"][0]["size_bytes"]>0
    assert OPS.visual_validate({"input_path":str(path),"max_pages":2})["summary"]["blank_sampled_pages"]==0


def test_docx_structure_tables_and_compare(tmp_path,monkeypatch):
    monkeypatch.setenv("AI_DATASEEK_OUTPUT_ROOT",str(tmp_path)); first=tmp_path/"first.docx"; second=tmp_path/"second.docx"; _docx(first); _docx(second)
    inspected=OPS.inspect_document({"input_path":str(first)}); assert inspected["summary"]["tables"]==1
    structure=OPS.docx_structure({"input_path":str(first),"max_chars":10000}); assert structure["summary"]["blocks"][0]["type"]=="heading"
    tables=OPS.docx_tables({"input_path":str(first),"output_dir":str(tmp_path/"tables"),"max_tables":5}); assert tables["summary"]["table_count"]==1
    compared=OPS.compare_documents({"left_path":str(first),"right_path":str(second),"max_chars":10000}); assert compared["summary"]["similarity"]==1.0


def test_pdf_ocr_is_page_bounded_and_writes_provenance(tmp_path, monkeypatch):
    from types import SimpleNamespace
    path=tmp_path/"scan.pdf"; _pdf(path)
    monkeypatch.setenv("AI_DATASEEK_OUTPUT_ROOT",str(tmp_path))
    monkeypatch.setattr(OPS.subprocess,"run",lambda *args,**kwargs:SimpleNamespace(returncode=0,stdout="温度 observations".encode(),stderr=b""))
    result=OPS.pdf_ocr({"input_path":str(path),"languages":"chi_sim+eng","max_pages":1,"output_path":str(tmp_path/"ocr.json")})
    assert result["summary"]["pages_ocrd"] == 1
    assert result["summary"]["pages"][0]["page"] == 1
    assert result["artifacts"][0]["size_bytes"] > 0


def test_pdf_ocr_rejects_invalid_languages(tmp_path):
    path=tmp_path/"scan.pdf"; _pdf(path)
    try: OPS.pdf_ocr({"input_path":str(path),"languages":"bad language"})
    except ValueError as exc: assert "language" in str(exc)
    else: raise AssertionError("unsafe OCR language was accepted")


def test_presentation_plugin_is_discoverable():
    import importlib.util
    plugin=Path(__file__).resolve().parents[2]/"tools"/"presentations"/"manifest.json"
    manifest=__import__("json").loads(plugin.read_text())
    assert {item["name"] for item in manifest["tools"]} == {"presentation_inspect","pptx_extract_content","pptx_render_slides"}
