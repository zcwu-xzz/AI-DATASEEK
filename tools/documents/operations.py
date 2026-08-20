#!/usr/bin/env python3
from __future__ import annotations

import argparse
import base64
import difflib
import json
import math
import os
import re
import subprocess
import tempfile
from pathlib import Path
from typing import Any

import numpy as np


def clean(value: Any) -> Any:
    if value is None or isinstance(value,(str,bool,int)): return value
    if isinstance(value,(float,np.floating)): return float(value) if math.isfinite(float(value)) else None
    if isinstance(value,np.integer): return int(value)
    if isinstance(value,dict): return {str(k):clean(v) for k,v in value.items()}
    if isinstance(value,(list,tuple,np.ndarray)): return [clean(v) for v in value]
    return str(value)


def response(operation: str, summary: dict[str,Any], *, evidence: list[dict[str,Any]]|None=None, artifacts: list[dict[str,Any]]|None=None, warnings: list[str]|None=None) -> dict[str,Any]:
    return clean({"success":True,"answer_ready":True,"operation":operation,"summary":summary,"evidence":evidence or [],"artifacts":artifacts or [],"warnings":warnings or [],"provenance":{"tool":operation,"version":"1.0.0"},"recommended_next_tools":[]})


def output_path(raw: str, directory: bool=False) -> Path:
    root=Path(os.environ.get("AI_DATASEEK_OUTPUT_ROOT","/home/ubuntu/output")).resolve(); path=Path(raw).resolve()
    if root!=path and root not in path.parents: raise ValueError(f"output path must be below {root}")
    (path if directory else path.parent).mkdir(parents=True,exist_ok=True); return path


def artifact(path: Path, mime: str) -> dict[str,Any]: return {"path":str(path),"type":mime,"size_bytes":path.stat().st_size}


def select_pages(document: Any, requested: list[int]|None, maximum: int) -> list[int]:
    if requested:
        indices=[]
        for page in requested:
            if page<1 or page>len(document): raise ValueError(f"page is outside document: {page}")
            if page-1 not in indices: indices.append(page-1)
        return indices[:maximum]
    if len(document)<=maximum: return list(range(len(document)))
    return sorted(set(np.linspace(0,len(document)-1,maximum,dtype=int).tolist()))


def document_text(path: str, max_chars: int) -> tuple[str,list[str]]:
    source,suffix=Path(path),Path(path).suffix.lower(); warnings=[]
    if suffix==".pdf":
        import pymupdf as fitz
        with fitz.open(source) as doc: text="\n\n".join(page.get_text("text") for page in doc)
        if not text.strip(): warnings.append("PDF has no extractable text layer and may require OCR")
    elif suffix==".docx":
        from docx import Document
        doc=Document(source); parts=[p.text for p in doc.paragraphs]; parts += ["\t".join(cell.text for cell in row.cells) for table in doc.tables for row in table.rows]; text="\n".join(parts)
    elif suffix==".doc":
        with tempfile.TemporaryDirectory() as temp:
            completed=subprocess.run(["libreoffice","--headless","--convert-to","txt:Text","--outdir",temp,str(source)],capture_output=True,text=True,timeout=90)
            converted=Path(temp)/(source.stem+".txt")
            if completed.returncode or not converted.exists(): raise ValueError("legacy DOC conversion failed")
            text=converted.read_text(encoding="utf-8",errors="replace")
    elif suffix in {".txt",".md"}: text=source.read_text(encoding="utf-8",errors="replace")
    else: raise ValueError("supported document formats are PDF, DOCX, DOC, TXT and Markdown")
    if len(text)>max_chars: warnings.append(f"text truncated at {max_chars} characters")
    return text[:max_chars],warnings


def inspect_document(a: dict[str,Any]) -> dict[str,Any]:
    path=Path(a["input_path"]); suffix=path.suffix.lower(); warnings=[]
    if suffix==".pdf":
        import pymupdf as fitz
        with fitz.open(path) as doc:
            metadata={k:v for k,v in doc.metadata.items() if v}; summary={"format":"pdf","pages":len(doc),"encrypted":bool(doc.needs_pass),"metadata":metadata,"text_pages":sum(bool(p.get_text().strip()) for p in doc),"image_count":sum(len(p.get_images(full=True)) for p in doc)}
            if summary["text_pages"]==0: warnings.append("PDF has no extractable text layer")
    elif suffix==".docx":
        from docx import Document
        doc=Document(path); props=doc.core_properties; summary={"format":"docx","paragraphs":len(doc.paragraphs),"tables":len(doc.tables),"inline_shapes":len(doc.inline_shapes),"sections":len(doc.sections),"metadata":{"title":props.title,"author":props.author,"subject":props.subject}}
    elif suffix==".doc":
        text,warnings=document_text(str(path),10000); summary={"format":"doc","legacy":True,"extractable_characters":len(text)}
    elif suffix in {".txt",".md"}:
        text=path.read_text(encoding="utf-8",errors="replace"); summary={"format":suffix[1:],"characters":len(text),"lines":text.count("\n")+1}
    else: raise ValueError("unsupported document format")
    summary.update(name=path.name,size_bytes=path.stat().st_size); return response("document_inspect",summary,warnings=warnings)


def pdf_text(a: dict[str,Any]) -> dict[str,Any]:
    import pymupdf as fitz
    limit=a.get("max_chars",200000); pages=[]; total=0; warnings=[]
    with fitz.open(a["input_path"]) as doc:
        for index in select_pages(doc,a.get("pages"),100):
            text=doc[index].get_text("text"); remaining=max(0,limit-total); text=text[:remaining]; pages.append({"page":index+1,"text":text}); total+=len(text)
            if total>=limit: warnings.append(f"text truncated at {limit} characters"); break
    artifacts=[]
    if a.get("output_path"):
        path=output_path(a["output_path"])
        if path.suffix.lower()==".txt": path.write_text("\n\n".join(f"[Page {p['page']}]\n{p['text']}" for p in pages),encoding="utf-8"); mime="text/plain"
        elif path.suffix.lower()==".json": path.write_text(json.dumps(pages,ensure_ascii=False,indent=2),encoding="utf-8"); mime="application/json"
        else: raise ValueError("PDF text output must be .txt or .json")
        artifacts=[artifact(path,mime)]
    return response("pdf_extract_text",{"pages_extracted":len(pages),"characters":total,"pages":pages},artifacts=artifacts,warnings=warnings)


def pdf_tables(a: dict[str,Any]) -> dict[str,Any]:
    import pymupdf as fitz
    import pandas as pd
    directory=output_path(a["output_dir"],directory=True); manifest=[]; artifacts=[]; maximum=a.get("max_tables",20)
    with fitz.open(a["input_path"]) as doc:
        for page_index in select_pages(doc,a.get("pages"),50):
            finder=doc[page_index].find_tables()
            for table_index,table in enumerate(finder.tables,1):
                rows=table.extract();
                if not rows: continue
                width=max(len(row) for row in rows); normalized=[list(row)+[None]*(width-len(row)) for row in rows]; frame=pd.DataFrame(normalized[1:],columns=[str(v or f"column_{i+1}") for i,v in enumerate(normalized[0])])
                path=directory/f"page-{page_index+1}-table-{table_index}.csv"; frame.to_csv(path,index=False); artifacts.append(artifact(path,"text/csv")); manifest.append({"page":page_index+1,"table":table_index,"rows":len(frame),"columns":list(frame.columns),"path":str(path),"bbox":list(table.bbox)})
                if len(manifest)>=maximum: break
            if len(manifest)>=maximum: break
    manifest_path=directory/"tables-manifest.json"; manifest_path.write_text(json.dumps(manifest,ensure_ascii=False,indent=2),encoding="utf-8"); artifacts.append(artifact(manifest_path,"application/json"))
    return response("pdf_extract_tables",{"table_count":len(manifest),"tables":manifest},artifacts=artifacts,warnings=[] if manifest else ["no native PDF tables were detected; scanned tables may require OCR"])


def pdf_render(a: dict[str,Any]) -> dict[str,Any]:
    import pymupdf as fitz
    directory=output_path(a["output_dir"],directory=True); artifacts=[]; rendered=[]; scale=a.get("dpi",144)/72
    with fitz.open(a["input_path"]) as doc:
        for index in select_pages(doc,a.get("pages"),a.get("max_pages",6)):
            pix=doc[index].get_pixmap(matrix=fitz.Matrix(scale,scale),alpha=False); path=directory/f"page-{index+1}.png"; pix.save(path); artifacts.append(artifact(path,"image/png")); rendered.append({"page":index+1,"width":pix.width,"height":pix.height,"path":str(path)})
    return response("pdf_render_pages",{"rendered":rendered},artifacts=artifacts)


def pdf_evidence(a: dict[str,Any]) -> dict[str,Any]:
    import pymupdf as fitz
    hits=[]; queries=[q for q in a["queries"] if q.strip()]; maximum=a.get("max_hits",50); context=a.get("context_chars",120)
    with fitz.open(a["input_path"]) as doc:
        for page_index,page in enumerate(doc):
            text=page.get_text("text"); lower=text.casefold()
            for query in queries:
                start=0
                while len(hits)<maximum:
                    found=lower.find(query.casefold(),start)
                    if found<0: break
                    hits.append({"query":query,"page":page_index+1,"snippet":text[max(0,found-context):found+len(query)+context].replace("\n"," ").strip(),"rectangles":[list(r) for r in page.search_for(query)[:10]]}); start=found+len(query)
            if len(hits)>=maximum: break
    return response("pdf_find_evidence",{"hit_count":len(hits),"queries":queries},evidence=hits,warnings=[] if hits else ["no matching extractable text was found"])


def pdf_ocr(a: dict[str,Any]) -> dict[str,Any]:
    import pymupdf as fitz
    languages=a.get("languages","chi_sim+eng")
    if not re.fullmatch(r"[A-Za-z0-9_+-]+",languages): raise ValueError("invalid OCR language selection")
    limit=a.get("max_chars",200000); maximum=a.get("max_pages",10); dpi=a.get("dpi",200); pages=[]; used=0; warnings=[]
    with fitz.open(a["input_path"]) as doc:
        indices=select_pages(doc,a.get("pages"),maximum)
        for index in indices:
            pix=doc[index].get_pixmap(matrix=fitz.Matrix(dpi/72,dpi/72),colorspace=fitz.csRGB,alpha=False)
            completed=subprocess.run(
                ["tesseract","stdin","stdout","-l",languages,"--dpi",str(dpi),"--psm","6"],
                input=pix.tobytes("png"),capture_output=True,timeout=90,
            )
            if completed.returncode:
                raise ValueError(f"OCR failed on page {index+1}: {completed.stderr.decode('utf-8',errors='replace')[:500]}")
            text=completed.stdout.decode("utf-8",errors="replace").strip(); remaining=max(0,limit-used); text=text[:remaining]
            pages.append({"page":index+1,"text":text,"characters":len(text)}); used+=len(text)
            if used>=limit: warnings.append(f"OCR text truncated at {limit} characters"); break
    artifacts=[]
    if a.get("output_path"):
        path=output_path(a["output_path"])
        if path.suffix.lower()==".txt": path.write_text("\n\n".join(f"[Page {p['page']}]\n{p['text']}" for p in pages),encoding="utf-8"); mime="text/plain"
        elif path.suffix.lower()==".json": path.write_text(json.dumps(pages,ensure_ascii=False,indent=2),encoding="utf-8"); mime="application/json"
        else: raise ValueError("PDF OCR output must be .txt or .json")
        artifacts=[artifact(path,mime)]
    if not any(page["text"] for page in pages): warnings.append("OCR produced no text on the selected pages")
    return response("pdf_ocr_text",{"pages_ocrd":len(pages),"characters":used,"languages":languages,"pages":pages},artifacts=artifacts,warnings=warnings)


def docx_structure(a: dict[str,Any]) -> dict[str,Any]:
    from docx import Document
    doc=Document(a["input_path"]); limit=a.get("max_chars",200000); used=0; blocks=[]; warnings=[]
    for paragraph in doc.paragraphs:
        text=paragraph.text
        if not text.strip(): continue
        remaining=limit-used
        if remaining<=0: warnings.append(f"content truncated at {limit} characters"); break
        style=paragraph.style.name if paragraph.style else None; blocks.append({"type":"heading" if style and style.lower().startswith("heading") else "paragraph","style":style,"text":text[:remaining]}); used+=len(text[:remaining])
    tables=[{"index":i+1,"rows":len(t.rows),"columns":max((len(r.cells) for r in t.rows),default=0),"preview":[[c.text for c in r.cells] for r in t.rows[:5]]} for i,t in enumerate(doc.tables)]
    summary={"blocks":blocks,"tables":tables,"image_count":len(doc.inline_shapes),"sections":len(doc.sections),"characters":used}; artifacts=[]
    if a.get("output_path"):
        path=output_path(a["output_path"])
        if path.suffix.lower()!=".json": raise ValueError("DOCX structure output must be .json")
        path.write_text(json.dumps(clean(summary),ensure_ascii=False,indent=2),encoding="utf-8"); artifacts=[artifact(path,"application/json")]
    return response("docx_extract_structure",summary,artifacts=artifacts,warnings=warnings)


def docx_tables(a: dict[str,Any]) -> dict[str,Any]:
    import pandas as pd
    from docx import Document
    directory=output_path(a["output_dir"],directory=True); doc=Document(a["input_path"]); manifest=[]; artifacts=[]
    for index,table in enumerate(doc.tables[:a.get("max_tables",30)],1):
        rows=[[cell.text for cell in row.cells] for row in table.rows]
        if not rows: continue
        width=max(len(row) for row in rows); rows=[row+[""]*(width-len(row)) for row in rows]; frame=pd.DataFrame(rows[1:],columns=[v or f"column_{i+1}" for i,v in enumerate(rows[0])]); path=directory/f"table-{index}.csv"; frame.to_csv(path,index=False); artifacts.append(artifact(path,"text/csv")); manifest.append({"table":index,"rows":len(frame),"columns":list(frame.columns),"path":str(path)})
    path=directory/"tables-manifest.json"; path.write_text(json.dumps(manifest,ensure_ascii=False,indent=2),encoding="utf-8"); artifacts.append(artifact(path,"application/json"))
    return response("docx_extract_tables",{"table_count":len(manifest),"tables":manifest},artifacts=artifacts)


def compare_documents(a: dict[str,Any]) -> dict[str,Any]:
    left,lw=document_text(a["left_path"],a.get("max_chars",300000)); right,rw=document_text(a["right_path"],a.get("max_chars",300000)); left_lines=[line.strip() for line in left.splitlines() if line.strip()]; right_lines=[line.strip() for line in right.splitlines() if line.strip()]; matcher=difflib.SequenceMatcher(None,left,right,autojunk=False); diff=list(difflib.unified_diff(left_lines,right_lines,fromfile=Path(a["left_path"]).name,tofile=Path(a["right_path"]).name,n=2)); summary={"similarity":matcher.ratio(),"left_characters":len(left),"right_characters":len(right),"diff_lines":diff[:2000],"diff_truncated":len(diff)>2000}; artifacts=[]
    if a.get("output_path"):
        path=output_path(a["output_path"])
        if path.suffix.lower()!=".json": raise ValueError("document comparison output must be .json")
        path.write_text(json.dumps(clean(summary),ensure_ascii=False,indent=2),encoding="utf-8"); artifacts=[artifact(path,"application/json")]
    return response("document_compare",summary,artifacts=artifacts,warnings=lw+rw)


def render_word_to_pdf(source: Path, directory: Path) -> Path:
    completed=subprocess.run(["libreoffice","--headless","--convert-to","pdf","--outdir",str(directory),str(source)],capture_output=True,text=True,timeout=90)
    result=directory/(source.stem+".pdf")
    if completed.returncode or not result.exists(): raise ValueError("LibreOffice could not render the Word document")
    return result


def visual_validate(a: dict[str,Any]) -> dict[str,Any]:
    import pymupdf as fitz
    source=Path(a["input_path"]); warnings=[]
    with tempfile.TemporaryDirectory() as temp:
        pdf=source if source.suffix.lower()==".pdf" else render_word_to_pdf(source,Path(temp))
        pages=[]
        with fitz.open(pdf) as doc:
            for index in select_pages(doc,None,a.get("max_pages",6)):
                page=doc[index]; pix=page.get_pixmap(matrix=fitz.Matrix(.5,.5),colorspace=fitz.csGRAY,alpha=False); values=np.frombuffer(pix.samples,dtype=np.uint8); standard=float(values.std()); text_chars=len(page.get_text("text")); blank=standard<1.0 and text_chars==0; pages.append({"page":index+1,"width":page.rect.width,"height":page.rect.height,"text_characters":text_chars,"pixel_standard_deviation":standard,"blank":blank})
                if blank: warnings.append(f"page {index+1} appears blank")
            summary={"pages":len(doc),"sampled_pages":pages,"blank_sampled_pages":sum(p["blank"] for p in pages),"rendered_from_word":source.suffix.lower()!=".pdf"}
    return response("document_visual_validate",summary,warnings=warnings)


FUNCTIONS={"document_inspect":inspect_document,"pdf_extract_text":pdf_text,"pdf_extract_tables":pdf_tables,"pdf_render_pages":pdf_render,"pdf_find_evidence":pdf_evidence,"pdf_ocr_text":pdf_ocr,"docx_extract_structure":docx_structure,"docx_extract_tables":docx_tables,"document_compare":compare_documents,"document_visual_validate":visual_validate}


def main() -> int:
    parser=argparse.ArgumentParser(); parser.add_argument("tool"); parser.add_argument("payload"); args=parser.parse_args()
    try: output=FUNCTIONS[args.tool](json.loads(base64.urlsafe_b64decode(args.payload+"="*(-len(args.payload)%4))))
    except Exception as exc: output={"success":False,"answer_ready":True,"operation":args.tool,"error":f"{type(exc).__name__}: {exc}","warnings":[]}
    print(json.dumps(clean(output),ensure_ascii=False,allow_nan=False)); return 0 if output["success"] else 1


if __name__=="__main__": raise SystemExit(main())
